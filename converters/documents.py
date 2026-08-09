import asyncio
import logging
import os
import urllib.request
import zipfile
import docx
import mammoth
import fitz  # PyMuPDF
from fpdf import FPDF
from pdf2docx import Converter
import markdown

logger = logging.getLogger(__name__)

def get_dejavu_font():
    font_path = "/tmp/DejaVuSans.ttf"
    if not os.path.exists(font_path):
        try:
            url = "https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans.ttf"
            urllib.request.urlretrieve(url, font_path)
        except Exception as e:
            logger.error(f"Failed to download DejaVu font: {e}")
    return font_path

def _docx_to_txt(input_path: str, output_path: str) -> str:
    try:
        doc = docx.Document(input_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        return output_path
    except Exception as e:
        logger.error(f"Error converting docx to txt: {e}")
        raise

async def convert_docx_to_txt(input_path: str, output_path: str) -> str:
    return await asyncio.to_thread(_docx_to_txt, input_path, output_path)

def _docx_to_html(input_path: str, output_path: str) -> str:
    try:
        with open(input_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html)
        return output_path
    except Exception as e:
        logger.error(f"Error converting docx to html: {e}")
        raise

async def convert_docx_to_html(input_path: str, output_path: str) -> str:
    return await asyncio.to_thread(_docx_to_html, input_path, output_path)

def _docx_to_pdf(input_path: str, output_path: str) -> str:
    try:
        doc = docx.Document(input_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        pdf = FPDF()
        pdf.add_page()
        font_path = get_dejavu_font()
        if os.path.exists(font_path):
            pdf.add_font("DejaVu", "", font_path, uni=True)
            pdf.set_font("DejaVu", size=12)
        else:
            pdf.set_font("Arial", size=12)
        
        pdf.multi_cell(0, 10, text)
        pdf.output(output_path)
        return output_path
    except Exception as e:
        logger.error(f"Error converting docx to pdf: {e}")
        raise

async def convert_docx_to_pdf(input_path: str, output_path: str) -> str:
    return await asyncio.to_thread(_docx_to_pdf, input_path, output_path)

def _pdf_to_txt(input_path: str, output_path: str) -> str:
    try:
        doc = fitz.open(input_path)
        text = ""
        for page in doc:
            text += page.get_text()
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        return output_path
    except Exception as e:
        logger.error(f"Error converting pdf to txt: {e}")
        raise

async def convert_pdf_to_txt(input_path: str, output_path: str) -> str:
    return await asyncio.to_thread(_pdf_to_txt, input_path, output_path)

def _pdf_to_png(input_path: str, output_path: str) -> str:
    try:
        doc = fitz.open(input_path)
        zip_path = output_path
        if not zip_path.endswith('.zip'):
            zip_path += '.zip'
        
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for i, page in enumerate(doc):
                pix = page.get_pixmap()
                img_path = f"/tmp/page_{i+1}.png"
                pix.save(img_path)
                zipf.write(img_path, f"page_{i+1}.png")
                os.remove(img_path)
        return zip_path
    except Exception as e:
        logger.error(f"Error converting pdf to png: {e}")
        raise

async def convert_pdf_to_png(input_path: str, output_path: str) -> str:
    return await asyncio.to_thread(_pdf_to_png, input_path, output_path)

def _txt_to_pdf(input_path: str, output_path: str) -> str:
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        pdf = FPDF()
        pdf.add_page()
        font_path = get_dejavu_font()
        if os.path.exists(font_path):
            pdf.add_font("DejaVu", "", font_path, uni=True)
            pdf.set_font("DejaVu", size=12)
        else:
            pdf.set_font("Arial", size=12)
        
        pdf.multi_cell(0, 10, text)
        pdf.output(output_path)
        return output_path
    except Exception as e:
        logger.error(f"Error converting txt to pdf: {e}")
        raise

async def convert_txt_to_pdf(input_path: str, output_path: str) -> str:
    return await asyncio.to_thread(_txt_to_pdf, input_path, output_path)

def _txt_to_docx(input_path: str, output_path: str) -> str:
    try:
        doc = docx.Document()
        with open(input_path, 'r', encoding='utf-8') as f:
            for line in f:
                doc.add_paragraph(line.strip())
        doc.save(output_path)
        return output_path
    except Exception as e:
        logger.error(f"Error converting txt to docx: {e}")
        raise

async def convert_txt_to_docx(input_path: str, output_path: str) -> str:
    return await asyncio.to_thread(_txt_to_docx, input_path, output_path)

def _pdf_to_docx(input_path: str, output_path: str) -> str:
    try:
        cv = Converter(input_path)
        cv.convert(output_path)
        cv.close()
        return output_path
    except Exception as e:
        logger.error(f"Error converting pdf to docx: {e}")
        raise

async def convert_pdf_to_docx(input_path: str, output_path: str) -> str:
    return await asyncio.to_thread(_pdf_to_docx, input_path, output_path)

def _md_to_html(input_path: str, output_path: str) -> str:
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()
        html = markdown.markdown(text)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
        return output_path
    except Exception as e:
        logger.error(f"Error converting md to html: {e}")
        raise

async def convert_md_to_html(input_path: str, output_path: str) -> str:
    return await asyncio.to_thread(_md_to_html, input_path, output_path)

def _md_to_pdf(input_path: str, output_path: str) -> str:
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()
        # Using basic text for PDF since fpdf2 HTML rendering might be complex, 
        # or just render text. For true HTML to PDF we might need other tools, 
        # but we'll stick to fpdf2 multi_cell.
        pdf = FPDF()
        pdf.add_page()
        font_path = get_dejavu_font()
        if os.path.exists(font_path):
            pdf.add_font("DejaVu", "", font_path, uni=True)
            pdf.set_font("DejaVu", size=12)
        else:
            pdf.set_font("Arial", size=12)
        
        pdf.multi_cell(0, 10, text)
        pdf.output(output_path)
        return output_path
    except Exception as e:
        logger.error(f"Error converting md to pdf: {e}")
        raise

async def convert_md_to_pdf(input_path: str, output_path: str) -> str:
    return await asyncio.to_thread(_md_to_pdf, input_path, output_path)

def _html_to_pdf(input_path: str, output_path: str) -> str:
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        pdf = FPDF()
        pdf.add_page()
        font_path = get_dejavu_font()
        if os.path.exists(font_path):
            pdf.add_font("DejaVu", "", font_path, uni=True)
            pdf.set_font("DejaVu", size=12)
        else:
            pdf.set_font("Arial", size=12)
        
        # Simple extraction
        from html.parser import HTMLParser
        class HTMLFilter(HTMLParser):
            text = ""
            def handle_data(self, data):
                self.text += data
        f = HTMLFilter()
        f.feed(text)
        
        pdf.multi_cell(0, 10, f.text)
        pdf.output(output_path)
        return output_path
    except Exception as e:
        logger.error(f"Error converting html to pdf: {e}")
        raise

async def convert_html_to_pdf(input_path: str, output_path: str) -> str:
    return await asyncio.to_thread(_html_to_pdf, input_path, output_path)


CONVERTERS = {
    ('docx', 'txt'): convert_docx_to_txt,
    ('docx', 'html'): convert_docx_to_html,
    ('docx', 'pdf'): convert_docx_to_pdf,
    ('pdf', 'txt'): convert_pdf_to_txt,
    ('pdf', 'png'): convert_pdf_to_png,
    ('txt', 'pdf'): convert_txt_to_pdf,
    ('txt', 'docx'): convert_txt_to_docx,
    ('pdf', 'docx'): convert_pdf_to_docx,
    ('md', 'html'): convert_md_to_html,
    ('md', 'pdf'): convert_md_to_pdf,
    ('html', 'pdf'): convert_html_to_pdf,
}
